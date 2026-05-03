from __future__ import annotations

import re
from pathlib import Path

from docx import Document


COMPANY_NAME = "Pacific Beacon Software, Inc."
COMPANY_PROFILE = (
    "Pacific Beacon Software, Inc. is a fictional mid-sized California software company "
    "with hybrid teams in San Francisco, Oakland, Sacramento, and several remote home offices. "
    "The company builds B2B workflow software and uses a standard support model across People Operations, "
    "Finance, Security, and a centralized IT Service Desk."
)
DOCUMENTS_DIR = Path(__file__).resolve().parent.parent / "documents"
EFFECTIVE_DATE = "March 1, 2026"
REVIEW_CADENCE = "Annual review or earlier if law, process, or business operations materially change."
MIN_WORDS = 2600


def theme(heading: str, topic: str, employee_action: str, california_note: str, manager_action: str, example: str) -> dict:
    return {
        "heading": heading,
        "topic": topic,
        "employee_action": employee_action,
        "california_note": california_note,
        "manager_action": manager_action,
        "example": example,
    }


def faq(question: str, answer: str) -> dict:
    return {"question": question, "answer": answer}


def doc(
    filename: str,
    title: str,
    owner: str,
    applies_to: str,
    purpose: str,
    audience_note: str,
    related: list[str],
    themes: list[dict],
    faqs: list[dict],
    contacts: list[str],
) -> dict:
    return {
        "filename": filename,
        "title": title,
        "owner": owner,
        "applies_to": applies_to,
        "purpose": purpose,
        "audience_note": audience_note,
        "related": related,
        "themes": themes,
        "faqs": faqs,
        "contacts": contacts,
    }


OPENERS = (
    "This section explains how the company expects the topic to work in normal day-to-day operations.",
    "The goal of this section is to reduce ambiguity so employees and managers can make the same decision from the same written standard.",
    "In practice, this topic matters because unclear guidance usually creates avoidable rework, escalation, or employee frustration.",
)
GOVERNANCE = (
    "Operationally, the company prefers a documented workflow and a visible decision trail over personal arrangements that only one manager remembers later.",
    "From an operating standpoint, the process is designed to be repeatable across departments so outcomes do not depend on personality or team habit.",
    "The company uses structured approvals and service channels here because small inconsistencies often turn into payroll, access, or employee-relations issues later.",
)
COLLAB = (
    "The collaboration expectation is that employees raise needs clearly, managers explain constraints specifically, and support teams document exceptions rather than guessing.",
    "Good outcomes depend on the employee, manager, and support function sharing the same facts at the same time instead of correcting the story after the decision was made.",
    "This topic works best when teams communicate early, document the result, and avoid side-channel decisions that are impossible to audit later.",
)


def render_theme_paragraphs(item: dict, index: int) -> list[str]:
    return [
        f"Pacific Beacon treats {item['topic']} as a normal operating topic that deserves explicit written guidance rather than informal habit. {item['topic'].capitalize()} affects employee trust, business continuity, and the quality of later decisions. {OPENERS[index % len(OPENERS)]}",
        f"Employees are expected to {item['employee_action']}. That expectation is intentionally practical: the company wants employees to know what good participation looks like before an issue becomes urgent or emotional. {GOVERNANCE[index % len(GOVERNANCE)]}",
        f"California context matters here because {item['california_note']}. For that reason, local custom or a quick verbal answer is not considered a substitute for the actual workflow, the official system of record, or the responsible owner.",
        f"Managers and support teams should {item['manager_action']}. The written process exists so similar situations get similar answers across departments, even when staffing, deadlines, or personalities vary. {COLLAB[index % len(COLLAB)]}",
        f"A typical operating example is this: {item['example']}. That example illustrates why the company documents the topic carefully: when time pressure rises, teams still need a predictable way to respond without relying on guesswork, side messages, or memory.",
    ]


def wrap_paragraph(text: str, width: int = 100) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current: list[str] = []
    current_len = 0
    for word in words:
        extra = len(word) + (1 if current else 0)
        if current and current_len + extra > width:
            lines.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len += extra
    if current:
        lines.append(" ".join(current))
    return lines


def wrap_bullets(items: list[str], kind: str, width: int = 96) -> list[str]:
    prefix = "- " if kind == ".md" else "* "
    lines: list[str] = []
    for item in items:
        wrapped = wrap_paragraph(item, width=width)
        if not wrapped:
            continue
        lines.append(prefix + wrapped[0])
        for continuation in wrapped[1:]:
            lines.append("  " + continuation)
    return lines


def render_text(spec: dict) -> str:
    kind = Path(spec["filename"]).suffix.lower()
    lines: list[str] = []

    def heading(text: str, level: int = 1) -> None:
        if kind == ".md":
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text.upper() if level == 1 else text)
        lines.append("")

    heading(COMPANY_NAME, 1)
    heading(spec["title"], 2 if kind == ".md" else 1)
    metadata = [
        ("Document Owner", spec["owner"]),
        ("Effective Date", EFFECTIVE_DATE),
        ("Review Cadence", REVIEW_CADENCE),
        ("Applies To", spec["applies_to"]),
        ("Purpose", spec["purpose"]),
        ("Audience Note", spec["audience_note"]),
        ("Related Documents", ", ".join(spec["related"])),
    ]
    for label, value in metadata:
        if kind == ".md":
            lines.append(f"**{label}:** {value}")
        else:
            lines.append(f"{label}: {value}")
    lines.append("")
    lines.extend(wrap_paragraph(COMPANY_PROFILE))
    lines.append("")
    lines.extend(wrap_paragraph(
        "This document is part of the internal demo knowledge base and is written to resemble a real handbook or operating guide that employees and managers could reference in a California company."
    ))
    lines.append("")

    heading("How to Use This Document", 2 if kind == ".md" else 2)
    lines.extend(wrap_paragraph(
        f"Use this document as the primary written guidance for {spec['title'].lower()}. Employees should read it for baseline expectations, while managers and support functions should use it as the shared reference point before creating one-off exceptions."
    ))
    lines.append("")
    lines.extend(wrap_paragraph(
        f"When a situation touches other company processes, reviewers should also consult {', '.join(spec['related'])} so decisions stay aligned across payroll, scheduling, access, onboarding, or benefits workflows."
    ))
    lines.append("")

    heading("Roles and Responsibilities", 2 if kind == ".md" else 2)
    lines.extend(wrap_paragraph(
        "Employees are responsible for raising requests early when possible, keeping information accurate, and using the designated workflow or service channel instead of relying solely on informal chat or private messages."
    ))
    lines.append("")
    lines.extend(wrap_paragraph(
        f"Managers are responsible for making timely, documented decisions, explaining operational constraints clearly, and partnering with {spec['owner']} whenever a case becomes sensitive, unusual, or legally significant."
    ))
    lines.append("")
    lines.extend(wrap_bullets([
        "Use the system of record for approvals, changes, and final status whenever one exists.",
        "Escalate early when policy interpretation, sensitive data, or employee-relations risk is involved.",
        "Keep explanations factual, professional, and specific enough to audit later.",
        "Avoid local side deals or verbal exceptions that contradict the written standard.",
    ], kind))
    lines.append("")

    for index, item in enumerate(spec["themes"]):
        heading(item["heading"], 2 if kind == ".md" else 2)
        for paragraph in render_theme_paragraphs(item, index):
            lines.extend(wrap_paragraph(paragraph))
            lines.append("")
        lines.extend(wrap_bullets([
            f"Employees should {item['employee_action']}.",
            f"Managers and support teams should {item['manager_action']}.",
            f"Escalate situations where {item['california_note']}.",
            "Document decisions in the normal workflow or system of record so the resolution can be reviewed later.",
        ], kind))
        lines.append("")

    heading("Records, Exceptions, and Escalation", 2 if kind == ".md" else 2)
    lines.extend(wrap_paragraph(
        "Exceptions should be rare, documented, and tied to a specific business reason. When a case cannot be resolved using the ordinary workflow, the reviewer should record the facts, identify the approving authority, and capture what follow-up is required."
    ))
    lines.append("")
    lines.extend(wrap_paragraph(
        f"The owning team for this document is {spec['owner']}. That team may partner with Finance, Security, IT, or People Operations depending on the issue, but the expectation is that the employee or manager receives a clear answer rather than being bounced between queues without ownership."
    ))
    lines.append("")
    lines.extend(wrap_bullets([
        "Record the date, decision, and owner when an exception is approved.",
        "Use a named escalation path for sensitive cases rather than informal forwarding.",
        "Correct the source system when possible instead of hiding the issue with a workaround.",
        "Close the loop with the employee after the issue is resolved.",
    ], kind))
    lines.append("")

    heading("Frequently Asked Questions", 2 if kind == ".md" else 2)
    for item in spec["faqs"]:
        if kind == ".md":
            lines.append(f"### {item['question']}")
        else:
            lines.append(f"Q: {item['question']}")
        lines.append("")
        lines.extend(wrap_paragraph(item["answer"]))
        lines.append("")

    heading("Support Contacts", 2 if kind == ".md" else 2)
    lines.extend(wrap_bullets(spec["contacts"], kind))
    lines.append("")

    heading("Revision Note", 2 if kind == ".md" else 2)
    lines.extend(wrap_paragraph(
        "This version was prepared for the demo environment to create a deeper, more realistic knowledge base. The language is intentionally practical so retrieval and summarization can surface policy-like answers instead of shallow FAQ snippets."
    ))
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def write_docx(path: Path, text: str) -> None:
    document = Document()
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("- ") or line.startswith("* "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue
        if line.startswith("**") and ":** " in line:
            label, value = line[2:].split(":** ", 1)
            paragraph = document.add_paragraph()
            paragraph.add_run(label + ": ").bold = True
            paragraph.add_run(value)
            continue
        document.add_paragraph(line)
    document.save(path)


def word_count(text: str) -> int:
    return len(re.findall(r"[A-Za-z0-9][A-Za-z0-9'/-]*", text))

DOCUMENT_SPECS = []

DOCUMENT_SPECS.extend([
    doc(
        "PTO_Policy.md",
        "California Leave and Time Away Policy",
        "People Operations",
        "All regular California employees, with manager guidance for approved contractors and interns.",
        "Define how employees plan vacation, use paid sick leave, request time away, and coordinate absences with managers, payroll, and team coverage.",
        "Primary employee-facing leave policy for the demo environment.",
        ["PTO_Vacation_Sick_Leave.docx", "Work_Schedules_and_Shifts.md", "Payroll_FAQ.md"],
        [
            theme("Policy Intent and Shared Expectations", "planned and unplanned leave", "request foreseeable time off early and review team calendars before submitting", "California paid sick leave, local ordinances, and protected leave rules can overlay the general policy, so edge cases should not be improvised", "use the HRIS workflow and make approval reasons visible instead of relying on memory", "a product manager wants three days off before a release cutoff and the manager confirms coverage in the system before approving"),
            theme("Accrual, Balance Use, and Carryover", "PTO balances and earned time", "review balances before booking travel or asking for extended time away", "earned time must be tracked carefully and discrepancies should be investigated through official records rather than personal spreadsheets", "encourage healthy time-off habits and flag balance anomalies quickly", "an employee notices a lower balance after a correction and opens a payroll-backed ticket with the relevant dates"),
            theme("Planned Vacation and Personal Time Requests", "vacation and personal time requests", "submit most requests at least ten business days in advance and include partial-day details when relevant", "advance notice supports planning, but managers should not use it to discourage lawful leave or reasonable requests", "evaluate actual coverage needs and propose alternatives when timing is difficult", "a customer success lead requests time off during quarter close and the manager approves with a backup owner"),
            theme("Sick Leave and Unexpected Absence", "illness, appointments, and same-day absence", "notify the manager or lead as soon as practical and share only the operational details needed for coverage", "medical privacy matters, and managers should not ask for diagnosis details beyond company guidance and legal requirements", "focus first on coverage, rescheduling, and work handoff rather than debate about the medical reason", "a support engineer wakes up ill before an early shift and the lead reassigns active tickets"),
            theme("Holidays, Closures, and Protected Leave Coordination", "company holidays and leave categories outside ordinary PTO", "check the holiday calendar and raise protected-leave needs early when possible", "protected leave administration, jury duty, bereavement, and related categories need formal handling instead of casual manager promises", "route possible protected-leave issues to People Operations and document temporary coverage changes", "an employee needs intermittent time for treatment and the manager handles short-term coverage while HR manages the leave path"),
            theme("Manager Review Standards, Payroll Impact, and Return to Work", "consistent approval decisions and accurate downstream records", "report pay or balance errors with dates, hours, and approval references", "California pay timing, wage statement detail, and final-pay sensitivity make accurate leave records especially important", "keep leave coding aligned across the HRIS, timekeeping, and team calendars and support re-entry after longer absences", "an employee returns from time away, checks changed priorities, and reports a missing holiday credit before payroll closes"),
        ],
        [
            faq("How far in advance should I request vacation?", "At least ten business days is the default expectation, but the better rule is early enough for real coverage planning and handoff preparation."),
            faq("Can I request half-day or hourly PTO?", "Yes. Partial-day requests are allowed when the time away affects work time, meetings, or payroll records."),
            faq("What if I get sick during a planned vacation?", "Contact your manager and People Operations promptly so the company can determine how the time should be coded under the applicable rules."),
            faq("Can my manager deny PTO?", "A manager may decline or move PTO when there is a concrete coverage, deadline, or operational issue, and the reason should be specific."),
            faq("Do company holidays reduce my PTO balance?", "Normally no. An observed company holiday inside an approved PTO window usually does not reduce your balance."),
            faq("Who should I contact if my balance looks wrong?", "Open a People Operations or payroll ticket with the dates involved and any supporting approval reference."),
            faq("What happens if I need an extended medical leave?", "Tell your manager you may need time away and contact People Operations quickly so formal leave administration can begin."),
            faq("Should I keep responding to messages while on PTO?", "The goal is real time away. Business-critical clarifications should be rare and limited."),
        ],
        ["People Operations Help Desk", "Payroll Support", "Assigned HR Business Partner"],
    ),
    doc(
        "PTO_Vacation_Sick_Leave.docx",
        "Manager Guide to Vacation and Sick Leave Administration",
        "People Operations Business Partners",
        "All California people managers and team leads who review absence requests.",
        "Provide a manager operating guide for reviewing, documenting, escalating, and communicating vacation, sick leave, and other short-term absences.",
        "Companion guide translating employee leave policy into manager actions.",
        ["PTO_Policy.md", "Work_Schedules.docx", "Payroll_FAQ.md"],
        [
            theme("How Managers Should Use This Guide", "leave administration as an operating discipline", "read the employee-facing policy first and then use this guide for manager steps", "leave, wage, and employee-relations risks often come from improvised manager behavior", "treat approvals, denials, and escalations as documented operating decisions", "a new manager resets old unwritten team customs to match the formal leave process"),
            theme("Evaluating Planned Leave Requests", "fair review of planned time off", "include enough context for the manager to understand timing, overlap, and coverage needs", "inconsistent denials or vague reasoning can create equity and compliance problems", "state the real business reason and discuss alternatives when possible", "two employees want the same holiday week and the manager uses documented coverage needs to decide"),
            theme("Handling Sick Leave and Unscheduled Absence", "same-day illness and unexpected absence", "notify quickly and provide the operational facts needed for immediate coverage", "managers should avoid informal documentation demands that conflict with paid sick leave rights", "reassign work first and route medical or extended-absence questions to People Operations", "a support representative calls out sick before a customer escalation shift and the manager immediately rotates ownership"),
            theme("Coverage Planning and Handoffs", "role backups and handoff standards", "follow the team handoff pattern and update shared calendars or work queues before leaving", "better coverage planning reduces the pressure to discourage lawful time off", "maintain backup lists, rotation plans, and realistic handoff expectations by role", "an engineering manager rotates release ownership before a high-vacation period"),
            theme("Escalation Triggers and Sensitive Cases", "cases that should not remain with the line manager alone", "raise accommodation, repeated-attendance, or protected-leave signals early instead of waiting for conflict", "California leave and retaliation concerns require early HR involvement when facts become sensitive", "share factual observations and business impact, then let People Operations guide the next step", "an employee has repeated illness-related absences after a prior flexibility request and the manager escalates facts rather than guessing intent"),
            theme("Communication, Documentation, and Return-to-Work Conversations", "clear written communication around absence decisions", "review the formal outcome, any schedule changes, and needed follow-up when returning to work", "return-to-work discussions should focus on workload and support, not protected medical details", "summarize decisions in the workflow and hold practical re-entry check-ins after longer absences", "an employee returns after two weeks away and the manager reviews changed priorities and unresolved system access"),
        ],
        [
            faq("Can I deny PTO because the team is busy?", "Only when there is a concrete coverage or timing issue. Document the reason and discuss alternatives when possible."),
            faq("Should I ask for a doctor note on day one?", "Not as a default. Follow company guidance and involve People Operations when the case becomes extended or sensitive."),
            faq("What if I already said yes in chat?", "Update the formal workflow and correct the record quickly. If circumstances changed, explain the issue transparently."),
            faq("How much handoff can I require?", "Only what is reasonably needed for continuity. Handoffs should fit the role and should not become hidden work during leave."),
            faq("What if two employees request the same holiday week?", "Use documented staffing needs, fairness over time, and prior approvals rather than arbitrary habit."),
            faq("When should I involve People Operations?", "Whenever protected leave, accommodations, repeated attendance patterns, or employee-relations risk is involved."),
            faq("Can I contact employees while they are out?", "Only for limited business-critical reasons or practical return-to-work coordination."),
            faq("What should I document after an absence?", "Document the dates, the coding path, major coverage decisions, and any payroll or HR follow-up needed."),
        ],
        ["People Operations Business Partner team", "HR Service Center", "Payroll escalation contact for leave coding"],
    ),
    doc(
        "Payroll_FAQ.md",
        "Payroll and Pay Practices Handbook",
        "Finance and Payroll Operations",
        "All U.S. employees, with California-specific notes called out where relevant.",
        "Explain payroll timing, direct deposit, payslips, deductions, reimbursements, final pay, and common pay-related questions in language employees can understand quickly.",
        "Employee-friendly payroll handbook for the demo.",
        ["Salary_and_Payroll.txt", "Benefits_Enrollment_Guide.md", "PTO_Policy.md"],
        [
            theme("Payroll Schedule and Payday Logic", "semi-monthly pay timing", "review payroll calendars and submit changes before cutoff", "California wage timing expectations are strict enough that missed deadlines can create real risk", "avoid promising same-cycle corrections until payroll confirms what is still possible", "the fifteenth falls on a Saturday and payroll moves the deposit date to the prior business day"),
            theme("Direct Deposit, Wage Statements, and Self-Service", "digital payroll delivery and self-service updates", "keep banking, tax, and address information accurate in the portal", "wage statement detail and accurate records matter because California pay disputes often depend on documentation quality", "remind employees to review payslips after compensation, leave, or benefits changes", "an employee changes banks just before close and payroll confirms whether the current cycle will still use the old account"),
            theme("Regular Earnings, Overtime, and Premium Pay", "base pay plus time-based earnings", "record hours worked accurately and report overtime, differentials, or premium situations promptly", "California overtime and premium-pay rules make precise timekeeping especially important", "review time records quickly and never encourage off-the-clock work", "a support analyst works late on an incident and the hours still need to be recorded and paid correctly"),
            theme("Deductions, Benefits, and Reimbursements", "how deductions and reimbursements interact with payroll", "compare benefit effective dates with the cycle in which deductions first appear and submit receipts promptly", "California business-expense reimbursement expectations reinforce the need for timely and documented claims", "coordinate with People Operations when a deduction issue starts in benefits data rather than payroll", "a medical election becomes effective on the first of the month and the employee sees the deduction on the first eligible payroll"),
            theme("Bonuses, Equity Events, and One-Time Payments", "non-routine payments", "read compensation letters carefully and ask whether one-time amounts are gross or net", "final-pay timing and documented approvals matter, so verbal promises are never enough for payroll action", "require the formal approval path before assigning a cycle to any one-time payment", "a manager says a spot bonus is approved but payroll waits for the signed finance workflow"),
            theme("Corrections, Final Pay, and Employee Questions", "pay corrections and separation timing", "report payroll issues quickly with the pay date, issue description, and supporting references", "California separation timing can create same-day or accelerated final-pay workflows", "take every pay concern seriously and route it for documented review", "an employee spots a missing overtime line on payday and submits the ticket immediately"),
        ],
        [
            faq("How often are salaries paid?", "Pacific Beacon pays on a semi-monthly schedule: normally the fifteenth and the last business day of the month."),
            faq("What happens if payday falls on a holiday?", "Payroll is generally issued on the prior business day so employees are not waiting for funds after a bank holiday or weekend."),
            faq("When should I update direct deposit information?", "As early as possible and before the published payroll cutoff. Late changes may miss the current cycle."),
            faq("Where can I find my payslips?", "In the employee self-service portal after payroll completes."),
            faq("How do reimbursements show up?", "Approved reimbursements usually appear in the next feasible payroll after final approval, but they are tracked separately from wages."),
            faq("What if my pay looks incorrect?", "Open a payroll ticket immediately with the pay date, issue description, expected outcome, and any supporting reference."),
            faq("Can payroll process a manager verbal promise?", "No. Payroll requires documented approvals for bonuses, compensation changes, and off-cycle items."),
            faq("Who handles final pay questions?", "Payroll and People Operations handle final-pay coordination together."),
        ],
        ["Payroll Operations inbox", "Finance shared services queue", "People Operations for leave-linked pay questions"],
    ),
    doc(
        "Salary_and_Payroll.txt",
        "Compensation and Payroll Operations Reference",
        "Finance Shared Services",
        "Finance, People Operations, and managers submitting pay-affecting changes.",
        "Serve as an operational companion to the employee payroll handbook by documenting approval gates, compensation change timing, and common payroll processing scenarios.",
        "Internal payroll operations reference rather than employee brochure text.",
        ["Payroll_FAQ.md", "Benefits_Enrollment_Guide.md", "PTO_Policy.md"],
        [
            theme("Compensation Governance and Source of Truth", "compensation controls and record ownership", "treat offer letters, approvals, HRIS records, and payroll entries as connected records", "California wage timing sensitivity makes sloppy late changes especially risky", "validate effective dates, approvals, and pay elements before processing", "a manager promises an increase immediately but payroll waits for the approved workflow and system update"),
            theme("Payroll Input Deadlines and Operational Cutoffs", "published close calendars and upstream dependencies", "submit timekeeping approvals, bonus files, and employment changes before internal cutoffs", "final-pay events and premium-pay corrections may justify exception handling when legally time sensitive", "communicate upstream changes before payroll close instead of treating payroll as a rescue team", "a promotion approved after salary close waits for the next run while a time-sensitive final-pay item jumps the queue"),
            theme("Off-Cycle Payroll and Supplemental Payments", "supplemental payroll as a limited exception path", "request off-cycle handling only when the business and legal facts justify it", "California wage timing issues are one of the clearest reasons a supplemental run may be necessary", "document the root cause and approving authority for each exception", "a missed commission file may or may not justify an off-cycle run depending on timing and employee impact"),
            theme("Benefits, Leave, and Timekeeping Intersections", "cross-functional issues that span multiple systems", "identify the system of origin before trying to fix a problem at the payment layer", "state-specific paid sick leave, reimbursement expectations, and wage statement detail make generic shortcuts unsafe", "align the factual sequence across payroll, People Operations, and the manager before communicating the answer", "an employee sick leave bank looks right but the payslip is wrong, so payroll checks the leave feed first"),
            theme("Auditability and Record Retention", "the ability to reconstruct what happened later", "keep approval references, change logs, and resolution notes tied to the actual transaction", "California disputes often turn on documentation quality, not just verbal recollection", "keep notes factual and close the loop with the employee once the correction is processed", "a delayed bonus is explained directly rather than with a vague payroll is reviewing message"),
            theme("Employee Communication and Ownership", "clear, respectful payroll communication", "submit tickets with dates, amounts, and the expected correction path", "pay questions are trust-sensitive and need visible ownership rather than queue bouncing", "name a responsible owner when multiple departments are involved", "a cross-functional issue is routed with a named owner and a final resolution note instead of three separate partial answers"),
        ],
        [
            faq("What is the best source for an employee-facing payroll answer?", "Start with Payroll_FAQ.md. Use this document when the issue becomes operational or control-oriented."),
            faq("When should payroll refuse a manual request?", "When it lacks approval, bypasses secure handling, or tries to avoid the HRIS and payroll control model."),
            faq("Can a manager request a same-day correction?", "Possibly, but Finance decides based on run status, legal requirements, and employee impact."),
            faq("Why do we ask for effective dates so often?", "Because compensation, benefits, and leave outcomes all depend on timing."),
            faq("Who owns a cross-functional pay issue?", "Ownership should be explicit. The team that identifies the root system should either resolve it or make a warm handoff."),
            faq("What belongs in an off-cycle request?", "Employee name, business reason, amount or pay element, urgency, approval reference, and root cause."),
            faq("How long should records be retained?", "Follow company retention standards and legal guidance, but the key principle is reconstructable payroll decisions."),
            faq("Why does explanation quality matter?", "Because payroll issues are personal and a precise explanation reduces confusion and repeat tickets."),
        ],
        ["Finance Shared Services manager", "Payroll Operations lead", "People Operations partner for compensation questions"],
    ),
])

DOCUMENT_SPECS.extend([
    doc(
        "Benefits_Enrollment_Guide.md",
        "Benefits Enrollment Guide",
        "People Operations Total Rewards",
        "Benefit-eligible employees in California and approved dependents where applicable.",
        "Guide employees through eligibility, enrollment windows, plan selection, dependents, life events, and payroll deduction timing with enough detail for self-service decisions.",
        "Practical open-enrollment handbook for a California software company.",
        ["Benefits.docx", "Payroll_FAQ.md", "PTO_Policy.md"],
        [
            theme("Eligibility and Enrollment Windows", "new-hire windows, open enrollment, and life-event changes", "review deadlines carefully and gather dependent documents early", "coverage continuation, leave, and payroll interactions can add California-specific timing questions", "point employees to Total Rewards before the window closes instead of after", "a new hire waits until the last day to gather marriage documentation and risks missing the enrollment cutoff"),
            theme("Medical, Dental, and Vision Plan Selection", "choosing health plans based on household needs", "compare provider access, payroll cost, and expected care usage before electing a plan", "California network availability can vary significantly by region and carrier", "explain plan mechanics in plain language without acting as a personal medical advisor", "an employee expecting regular specialist visits chooses broader access even with a higher payroll deduction"),
            theme("HSA, FSA, and Tax-Advantaged Accounts", "tax-advantaged accounts tied to benefits elections", "confirm plan compatibility and estimate contributions conservatively", "state tax treatment may differ from federal treatment, so assumptions should be checked", "explain payroll setup while leaving personal tax advice to the employee advisor", "an employee elects an FSA for predictable prescription costs and locks the amount before open enrollment closes"),
            theme("Life Insurance, Disability, and Income Protection", "core protection benefits", "review beneficiary designations and optional elections at hire and after major life changes", "state disability and paid family leave coordination can interact with company-sponsored coverage", "direct case-specific disability questions to Total Rewards rather than guessing", "an employee welcoming a child reviews beneficiary updates, life insurance, and paid leave coordination together"),
            theme("Dependents and Qualifying Life Events", "mid-year changes controlled by valid events", "contact Total Rewards early when marriage, birth, adoption, divorce, or loss of other coverage happens", "domestic partner and household-coverage questions may require California-specific review", "verify documentation timing and event dates before promising a change", "an employee loses other coverage near month-end and alerts Total Rewards immediately to preserve the window"),
            theme("Payroll Deductions, Claims, and Ongoing Maintenance", "what happens after enrollment is submitted", "review the first payroll cycles after enrollment and keep dependent and beneficiary data current", "leave, reimbursement, and wage statement questions can touch several systems at once", "route carrier questions, deduction questions, and life-event corrections to the right owner with dates and screenshots", "a deduction appears in payroll while the carrier site still shows no dependent and Total Rewards checks the carrier feed"),
        ],
        [
            faq("When can I enroll in benefits?", "Normally as a new hire, during annual open enrollment, or after an approved qualifying life event."),
            faq("When do my deductions start?", "Usually in the first payroll cycle after the plan becomes effective and the enrollment data reaches payroll."),
            faq("Can I add a dependent later?", "Only if you are still in a valid window or you experience a qualifying life event and provide documentation on time."),
            faq("How do I know which medical plan is best?", "Compare provider access, paycheck cost, expected care use, deductible exposure, and account compatibility."),
            faq("What if the carrier cannot find my enrollment?", "Contact Total Rewards with the submission date, confirmation information, and the specific coverage issue."),
            faq("Can payroll explain my HSA tax impact?", "Payroll can explain deduction setup, but personal tax advice should come from your own advisor."),
            faq("Do I need to update beneficiaries after life changes?", "Yes. Marriage, divorce, birth, adoption, and similar events are strong reasons to review beneficiary designations."),
            faq("Who helps with leave-related benefit questions?", "Start with Total Rewards or People Operations because leave coordination often spans payroll, carriers, and policy rules."),
        ],
        ["Total Rewards support inbox", "Benefits carrier service center", "Payroll support for deduction timing questions"],
    ),
    doc(
        "Benefits.docx",
        "Total Rewards and Benefits Overview",
        "People Operations Total Rewards",
        "All benefit-eligible employees and managers supporting benefit questions.",
        "Provide a broader benefits and perks narrative that complements the enrollment guide by explaining why each program exists and how employees typically use it.",
        "Brochure-style companion to the enrollment guide.",
        ["Benefits_Enrollment_Guide.md", "PTO_Policy.md", "Work_Schedules_and_Shifts.md"],
        [
            theme("Benefits Philosophy and Employee Experience", "benefits as part of the full employment experience", "treat benefits as a year-round resource rather than a once-a-year form", "California hiring markets and living costs make practical, well-explained benefits especially important", "highlight resources and enrollment steps without improvising technical advice", "a manager points a new hire to the portal, the EAP, and the open questions they should resolve before the deadline"),
            theme("Healthcare Coverage and Everyday Use", "using healthcare coverage effectively after enrollment", "activate carrier accounts, confirm provider access, and review routine care options before an urgent need appears", "network access can vary across California regions and family locations", "encourage employees to compare how they will actually use the plan, not just the deduction amount", "an employee with frequent pediatric visits chooses a plan based on provider access rather than premium alone"),
            theme("Mental Health, Family Support, and EAP Resources", "daily-life support beyond major medical events", "review counseling, referral, and work-life resources early rather than waiting for a crisis", "commuting, caregiving, and housing pressures can make preventive support especially valuable for California employees", "normalize use of support resources without requesting private details", "a manager reminds a stretched team about EAP access during a demanding release cycle"),
            theme("Retirement and Financial Wellness", "long-term planning and payroll-linked savings", "review retirement deferrals and beneficiary updates when compensation or household needs change", "high-cost markets make consistent savings habits especially important even when contribution amounts begin modestly", "remind employees that payroll setup is only one part of financial planning", "after a salary increase, an employee raises retirement contributions instead of letting the extra amount disappear into routine spending"),
            theme("Perks, Learning, and Hybrid Work Support", "supplemental programs that support hybrid work and professional growth", "read the current program guide before buying equipment, requesting reimbursement, or using a learning budget", "California hybrid work patterns often create practical home-office and commuting questions that need clear rules", "apply perk rules consistently and avoid side promises outside the published standard", "an employee wants a new monitor and confirms whether it should be reimbursed, issued as an asset, or pre-approved"),
            theme("Manager Role in Benefits Awareness", "how managers support understanding without acting as benefits specialists", "ask process questions early and use the official program contacts for detailed answers", "benefit misunderstandings can create unnecessary stress when life events, leave, or deductions change at the same time", "help employees find the right owner instead of guessing about plan details", "a manager helps an employee locate the right disability and leave contacts during a family event"),
        ],
        [
            faq("What is the difference between this guide and the enrollment guide?", "This guide explains the value and use of the programs, while the enrollment guide explains windows, elections, and mechanics."),
            faq("Can my manager tell me which plan to choose?", "Managers may share process guidance but should not act as personal medical, legal, or tax advisors."),
            faq("Where do I go for mental health support?", "Start with the plan materials or EAP resources, and use Total Rewards if you need help finding the correct path."),
            faq("How often should I review my benefits?", "At minimum during open enrollment and after major household, financial, or family-status changes."),
            faq("Are perks the same for everyone?", "Some are broad-based while others depend on role, location, or business need."),
            faq("How do home-office benefits interact with IT equipment?", "Some items are reimbursed, some are issued as company assets, and some require pre-approval."),
            faq("Can payroll answer all benefit questions?", "Payroll helps with deduction timing, but carrier, coverage, and life-event questions usually belong with Total Rewards."),
            faq("Why does the company emphasize benefits education?", "Because employees use programs better when they understand them in plain language rather than treating them as opaque forms."),
        ],
        ["Total Rewards team", "Learning and development program owner", "IT asset support for hybrid equipment questions"],
    ),
    doc(
        "Work_Schedules_and_Shifts.md",
        "Work Schedules, Hybrid Work, and Shift Coverage Policy",
        "People Operations and Department Managers",
        "All employees, with additional rules for non-exempt and shift-based roles.",
        "Define working hours, hybrid expectations, timekeeping, overtime approvals, meal and rest breaks, shift coverage, and schedule changes.",
        "Core employee-facing schedule policy for the demo.",
        ["Work_Schedules.docx", "PTO_Policy.md", "IT_Support_and_Access.md"],
        [
            theme("Standard Hours, Core Collaboration Time, and Local Flexibility", "shared hours plus department-level flexibility", "learn the default team schedule and request alternate start times or temporary changes through the approved path", "California meal, rest-break, overtime, and recordkeeping rules still apply even when teams are flexible", "publish team norms clearly and review recurring exceptions for fairness and business fit", "an engineer adjusts start time for school drop-off but still joins core design-review hours"),
            theme("Hybrid Work and On-Site Presence Expectations", "hybrid work as a structured operating model", "raise location changes early when badge access, equipment, or meeting support may be affected", "commuting, ergonomic, and reimbursable-expense realities make California hybrid work more detailed than simple work-from-home permission", "define anchor days and role-based on-site needs without leaving them as vague culture expectations", "a recruiting coordinator works remotely most days but is expected on site for interview loops"),
            theme("Timekeeping, Overtime, and Extra-Hour Approval", "recording time and handling extra hours", "record hours actually worked and report extra time or differential situations accurately", "California overtime and premium rules make precise timekeeping essential for non-exempt roles", "approve extra hours where possible but never pressure employees to under-report time", "a production incident runs late and the hours are still recorded even though the overtime was not ideal from a planning standpoint"),
            theme("Meal Periods, Rest Breaks, and Shift Practicalities", "break compliance as part of schedule design", "tell the manager when workload repeatedly prevents normal breaks", "meal and rest-break compliance is a real operating risk, not only a theoretical legal issue", "build staffing and shift patterns that make breaks realistic for queue-based and customer-facing roles", "a support queue spikes during lunch and the supervisor rotates coverage instead of normalizing missed breaks"),
            theme("Shift Swaps, On-Call, and Temporary Changes", "controlled flexibility in coverage-heavy teams", "treat a coworker agreement as incomplete until the manager approves the actual schedule change", "California overtime and break consequences can change when actual hours move around", "confirm accountability, handoff expectations, and pay consequences before a swap becomes final", "two analysts agree to swap Friday and Monday coverage and the supervisor validates the hours and records"),
            theme("Attendance, Communication, and Emergency Closures", "how employees respond to lateness, absence, or disruption", "use the fastest practical notification path and explain the operational impact when attendance changes", "California commute, safety, and pay questions can interact when office closures or home-office failures happen", "use centralized closure notices and clear escalation paths instead of leaving employees to guess", "a home power outage prevents remote work and the employee immediately confirms whether alternate work or time away is appropriate"),
        ],
        [
            faq("Can I change my regular working hours?", "Possibly. Changes depend on role needs, team coverage, manager approval, and whether the schedule still supports collaboration and compliance."),
            faq("Do remote employees still have core hours?", "Yes. Hybrid or remote work changes location, not the need for shared collaboration windows."),
            faq("What if I work extra hours without prior approval?", "You should still record the time actually worked. Pay accuracy comes first."),
            faq("Can coworkers swap shifts on their own?", "No. A swap is not final until the manager approves it and the records are updated."),
            faq("Who decides whether I need to be on site?", "Your department and manager define role-based on-site expectations and planned events."),
            faq("What if I miss a meal period because of workload?", "Tell your manager promptly so the issue can be reviewed and corrected."),
            faq("How should I report being late or unexpectedly absent?", "Use the team fastest practical notification path and provide enough context for coverage."),
            faq("Does hybrid work mean I can work from anywhere?", "No. Approved work locations still need to meet security, privacy, and operational expectations."),
        ],
        ["People Operations service center", "Department manager or scheduling lead", "Payroll support for timekeeping-linked pay questions"],
    ),
    doc(
        "Work_Schedules.docx",
        "Manager Scheduling and Coverage Playbook",
        "People Operations Business Partners",
        "Managers and scheduling leads responsible for staffing, overtime approval, and schedule communications.",
        "Translate the employee-facing schedule policy into manager practices for staffing, coverage planning, break compliance, and response to schedule exceptions.",
        "Playbook version of the schedule policy for supervisors.",
        ["Work_Schedules_and_Shifts.md", "PTO_Policy.md", "IT_Onboarding_Guide.md"],
        [
            theme("Staffing Models and Role-Based Scheduling", "designing schedules from the actual service model", "learn the baseline team coverage logic so requests are evaluated in the right context", "California break and pay rules still constrain schedule design even in flexible teams", "make the business model visible instead of letting schedules look arbitrary", "a support manager staggers start times to cover east-coast tickets without forcing the entire team into one pattern"),
            theme("Approving Flexibility Without Losing Control", "structured schedule flexibility", "treat recurring flexibility as a documented arrangement rather than an informal favor", "informal flexibility can create hidden overtime and break issues for non-exempt roles", "define whether an arrangement is temporary, ongoing, or experimental and when it will be reviewed", "a team member receives an earlier start time for caregiving reasons on a trial basis"),
            theme("Managing Overtime, Premium Risk, and Break Compliance", "the manager role in payroll risk", "surface repeated late work or missed breaks rather than assuming the team will absorb them quietly", "California premium and overtime obligations make staffing discipline essential", "treat recurring extra hours as a capacity or planning problem that needs escalation", "a finance close cycle regularly runs late and the manager adjusts staffing instead of living on overtime"),
            theme("Coverage Planning for PTO, Illness, and Business Peaks", "backup ownership and minimum coverage levels", "follow the coverage pattern and handoff expectations before peak periods begin", "better planning reduces the temptation to discourage leave or hide schedule pressure inside informal norms", "publish who owns approvals, customer escalations, and critical decisions before absences happen", "before a holiday week the manager confirms backup ownership for customer escalations and finance approvals"),
            theme("Communication Templates and Escalation", "clear written scheduling communication", "read and keep the written schedule decisions rather than relying on spoken assumptions", "schedule messages can create compliance problems when pay, breaks, or leave rights are affected", "confirm decisions in writing and escalate unusual or contested cases while there is still time to resolve them", "a manager documents a denial with the actual coverage reason and a proposed alternative"),
            theme("Schedule Audits and Equity Review", "using data instead of habit to review schedule patterns", "raise fairness questions when the same people always receive preferred schedules or always absorb overtime", "equity issues can become morale and legal problems if no one reviews them intentionally", "audit patterns quarterly and use the results to improve staffing rather than only documenting the imbalance", "a manager notices the same employees keep absorbing late coverage and rebalances future assignments"),
        ],
        [
            faq("What is the manager first responsibility in scheduling?", "Make the business model visible and staff it deliberately. Random scheduling decisions create morale and compliance risk."),
            faq("Can I allow flexibility informally if the employee is reliable?", "Small flexibility choices still need documented expectations so they remain fair and operationally sustainable."),
            faq("How do I know a schedule issue needs escalation?", "Escalate when pay treatment, repeated overtime, break compliance, leave sensitivity, or employee-relations tension is involved."),
            faq("Should I always deny PTO during peak periods?", "No. Use actual minimum coverage needs and planning tools before defaulting to blanket denials."),
            faq("How often should I review my team schedule patterns?", "Quarterly is a strong baseline, and sooner if staffing or repeated exceptions suggest drift."),
            faq("What if employees solve coverage themselves?", "Confirm the arrangement, because the manager remains accountable for the final schedule and pay consequences."),
            faq("How should I document a denied request?", "State the concrete operational reason, the decision date, and any alternative or follow-up plan."),
            faq("Why audit schedule equity?", "Because the same people often absorb overtime or lose preferred time off when leaders rely on habit instead of data."),
        ],
        ["People Operations business partner", "Department scheduling lead", "Payroll support for schedule-to-pay questions"],
    ),
])

DOCUMENT_SPECS.extend([
    doc(
        "IT_Support_and_Access.md",
        "IT Support and Access Guide",
        "IT Service Desk and Security Operations",
        "All employees, approved contingent workers, and managers submitting access requests.",
        "Explain how employees get help with hardware, software, password resets, VPN, system access, and security-sensitive issues.",
        "Core employee-facing IT support guide for the demo.",
        ["IT_Support.txt", "Access_Policy.md", "IT_Onboarding_Guide.md"],
        [
            theme("Support Channels, Coverage Model, and Ticket Priorities", "the Service Desk as the front door to support", "open tickets with the affected system, exact symptom, and business impact", "hybrid work across California offices and homes means remote context often changes the support path", "triage by business impact, incident scope, and security sensitivity instead of job title alone", "a laptop issue blocks production access and is handled differently from a routine accessory request"),
            theme("Identity, Password Reset, and MFA", "common identity recovery needs", "use self-service recovery first and update MFA methods before losing old devices when possible", "remote work means recovery often happens away from an office, so identity proofing still needs to remain strong", "never point employees toward shared credentials or unofficial shortcuts", "an employee is locked out after changing phones and follows recovery steps before opening a manual reset case"),
            theme("VPN, Remote Connectivity, and Secure Off-Network Work", "remote access to internal systems", "test required VPN access before critical meetings, travel-heavy weeks, or close periods", "California employees often work from homes, client sites, and temporary locations, which makes secure remote readiness a normal operating expectation", "encourage teams to validate remote access in advance rather than discovering a problem minutes before a deadline", "a finance analyst tests VPN and ERP access the day before remote close work"),
            theme("Software Requests, Role-Based Access, and Approvals", "access requests tied to business need", "submit software and access requests through the approved workflow with a clear business reason", "remote and hybrid work reinforce the need for least-privilege access and visible approval paths", "start with standard role profiles and use broader access only when the task truly requires it", "a customer success manager requests the standard CRM reporting profile rather than broad administrator rights"),
            theme("Hardware, Asset Lifecycle, and Remote Device Support", "company devices as managed assets", "use issued devices and report hardware issues quickly instead of improvising with unapproved personal tools", "remote and home-office work make shipping, repair, and replacement logistics part of normal support", "prioritize accelerated replacement only when business impact justifies it", "a sales employee with a failed webcam before a roadshow receives a prioritized accessory replacement"),
            theme("Security Incidents, Suspicious Messages, and Escalation Rules", "support issues that overlap with security operations", "report suspicious emails, repeated MFA prompts, lost devices, and unusual logins immediately", "California operations may involve customer, employee, and financial data that heighten the impact of delayed reporting", "support rapid escalation and avoid delaying reports until after a meeting or release", "an employee receives repeated unrecognized MFA prompts and reports them before retrying access"),
        ],
        [
            faq("How do I request VPN access?", "Submit an access request with the business reason and any required manager approval."),
            faq("What should I do if I forget my password?", "Use self-service password reset first. If that fails, contact the Service Desk for identity-verified recovery."),
            faq("Can I ask a coworker to share a login temporarily?", "No. Shared credentials are prohibited and create security and audit problems."),
            faq("How are software requests approved?", "Approval depends on role need, licensing, data sensitivity, and the system owner standard."),
            faq("What counts as a high-priority IT issue?", "An issue is high priority when it blocks work, affects multiple users, impacts a customer event, or creates a security risk."),
            faq("Do I need VPN from home?", "For many internal systems, yes. Follow the system guidance and use approved secure methods."),
            faq("What should I do if my laptop is lost?", "Report it immediately so remote protection, tracking, and replacement steps can begin."),
            faq("Who handles suspicious MFA prompts or phishing?", "Report them right away. The Service Desk or Security team will guide the next step."),
        ],
        ["IT Service Desk", "Security Operations escalation path", "IT asset operations team"],
    ),
    doc(
        "IT_Support.txt",
        "Service Desk Operations Reference",
        "IT Service Desk",
        "Service Desk analysts, local IT coordinators, and managers routing employee support questions.",
        "Provide the internal operational companion to the employee IT guide, with triage expectations, escalation cues, and common support patterns.",
        "Internal help desk runbook rather than employee-facing copy.",
        ["IT_Support_and_Access.md", "Access_Policy.md", "IT_Onboarding_Guide.md"],
        [
            theme("Triage Standards and First-Response Quality", "the quality of the first ticket note", "include the service affected, the symptom, the impact, and the current user context", "remote and multi-office California work patterns mean on-site versus home-office context can materially change the path", "categorize incidents and requests carefully so downstream queues are not forced to guess", "a vague VPN is broken note becomes much more useful when the analyst captures network type and exact symptom"),
            theme("Identity Proofing and Assisted Recovery", "manual help with passwords and MFA", "follow the identity-proofing standard before any assisted recovery action", "remote employees are common, so proofing must work even when the user is not at a desk in an office", "treat manager confirmation as urgency context, not as a substitute for identity verification", "an employee calls from a new phone number after travel and still completes the approved recovery checks"),
            theme("Access Request Routing and Least-Privilege Review", "reviewing access tickets for business fit", "describe the business task so the desk can match the right role profile instead of a vague broad request", "remote and hybrid work make over-broad access harder to notice if no one reviews it deliberately", "ask clarifying questions when a request sounds wider than the task described", "a manager asks for full finance access but the analyst only needs read-only reporting and invoice review"),
            theme("Asset Issues, Shipping, and Remote Replacement", "device support as an end-to-end process", "state whether work is blocked, whether any backup exists, and whether shipping is feasible", "California home-office and office-hopping patterns make shipping and loaner logistics part of ordinary support", "track asset status, serials, and return expectations instead of treating replacement as a one-message favor", "a field employee with a damaged laptop two days before travel needs a fast replacement path with proper tracking"),
            theme("Escalation Matrix, Outages, and Communications", "broad incidents and structured updates", "watch for pattern signals and update the user-facing message path instead of treating every ticket as isolated", "remote work can make a shared outage look like unrelated individual cases at first", "use concise, time-stamped status language that managers can relay without improvisation", "multiple users report authentication trouble after an identity-service change and the desk flags a wider incident"),
            theme("Case Notes, Ownership, and Closure Quality", "the long-term value of disciplined ticket notes", "record the routing logic, actions taken, and final root cause clearly", "documentation quality matters when incidents later become audit, security, or employee-experience questions", "close the loop only after the employee confirms usable working state or the root cause is communicated", "a case closes with a full note that lets the next analyst understand exactly what failed and what changed"),
        ],
        [
            faq("What belongs in the first ticket note?", "The affected service, the symptom, the business impact, and the current context such as remote or on-site status."),
            faq("When should I escalate a ticket to Security?", "When credential compromise, phishing, suspicious MFA prompts, unauthorized access, or lost devices are involved."),
            faq("Do managers count as identity proof?", "No. Managers can confirm urgency or employment context, but the user identity still needs approved verification."),
            faq("How broad should access requests be granted?", "Grant only what the role needs and ask clarifying questions when a request sounds too broad."),
            faq("What is a good outage communication?", "It states what is affected, what users should do now, when the next update is expected, and where official status will appear."),
            faq("When do I open a second ticket?", "Only when the issue is clearly different. Duplicate tickets usually slow triage."),
            faq("What if a remote device replacement is urgent?", "Capture the business deadline, current working state, and shipping constraints so asset operations can prioritize appropriately."),
            faq("Why do case notes matter so much?", "Because the next analyst, escalation team, or auditor should be able to understand what happened without relying on memory."),
        ],
        ["IT Service Desk lead", "Identity and access management queue", "Asset operations coordinator"],
    ),
    doc(
        "IT_Onboarding_Guide.md",
        "New Hire IT Onboarding Guide",
        "IT Operations and People Operations",
        "Hiring managers, recruiters, People Operations coordinators, and new hires receiving company systems.",
        "Explain how hardware, accounts, security setup, and role-based applications are prepared before and after a new hire starts.",
        "Cross-functional onboarding runbook for a California tech company.",
        ["IT_Support_and_Access.md", "Access_Policy.md", "Work_Schedules_and_Shifts.md"],
        [
            theme("Preboarding Timeline and Data Readiness", "what must be ready before day one", "confirm start date, role, manager, location, and shipping details early enough for staging", "California home-office shipping and office logistics often require more lead time than teams expect", "avoid last-minute onboarding data changes when hardware and access setup are already in motion", "a recruiter changes a start date after hardware ships and the coordinator must realign activation timing"),
            theme("Day One Setup: Device, Identity, and Collaboration Tools", "the first operational setup experience", "complete login, MFA, password setup, email, chat, and meeting access before worrying about advanced tools", "remote California hires may complete the whole sequence from a home office without desk-side support", "schedule enough Day One time for setup instead of filling the calendar with back-to-back meetings", "a remote engineer receives the laptop early but Day One activation still happens on the formal start date"),
            theme("Role-Based Application Bundles and Access Expansion", "baseline access plus later additions", "use the standard role bundle first and request extra access only when the actual task requires it", "least-privilege access matters even more when new hires begin from remote locations", "separate must-have Day One access from items that can wait until training is complete", "a support new hire receives ticketing and telephony access first while advanced reporting rights come later"),
            theme("Security Training, Data Handling, and Acceptable Use", "security as part of onboarding instead of a later add-on", "complete early security training and use approved tools from the beginning", "customer, employee, and financial data handled in California operations make early secure habits especially important", "reinforce that productivity never justifies credential sharing or unapproved storage", "a new employee wants to forward documents to a personal email and is redirected to approved tools"),
            theme("Manager Responsibilities During the First Two Weeks", "the manager role in successful onboarding", "review actual working access with the new hire once real team tasks begin", "hybrid onboarding can feel fragmented, so planned check-ins matter more than silence", "use first-week checkpoints to catch true blockers before frustration grows", "a new analyst can log in but cannot access shared finance folders and the manager catches the gap in a planned review"),
            theme("Support Paths for Late Hardware, Missing Access, and Setup Failures", "what happens when day-one setup does not go smoothly", "report blockers quickly and include the exact missing capability", "remote and home-office starts can make shipping, authentication, or local connectivity issues more likely", "decide explicitly whether the answer is a backup device, staged access, or a revised first-day plan", "hardware is delayed and IT, People Operations, and the manager agree on a temporary plan instead of leaving the new hire guessing"),
        ],
        [
            faq("When should IT start preparing for a new hire?", "As soon as the core employment data is confirmed and the start date is stable enough to stage hardware and identity setup."),
            faq("What should be working on Day One?", "At minimum the employee should be able to log in, use MFA, access email and messaging, join meetings, and reach baseline systems for the role."),
            faq("Can all applications be provisioned before the start date?", "Some can, but access activation timing is controlled so identity, security, and employment status stay aligned."),
            faq("Who decides which role bundle a new hire receives?", "IT uses the role, department, and manager-approved baseline profile, then routes extra requests through the normal access process."),
            faq("What if hardware does not arrive on time?", "IT and the manager should quickly decide whether a backup device, shipment acceleration, or revised start-day plan is needed."),
            faq("Why do managers need a first-week access check?", "Because many gaps appear only after the employee begins real team tasks."),
            faq("Should new hires use personal devices if the company device is delayed?", "Only if an approved exception exists. In most cases the safer path is to wait for managed equipment or IT-approved alternatives."),
            faq("How does security fit into onboarding?", "Security is embedded from the start through MFA, approved tools, acceptable use, and reporting expectations."),
        ],
        ["IT onboarding coordinator", "People Operations onboarding specialist", "Hiring manager escalation path"],
    ),
    doc(
        "Access_Policy.md",
        "Identity, Access, and Acceptable Use Policy",
        "Security Operations and Identity Governance",
        "All employees, approved contractors, system owners, and managers requesting or approving access.",
        "Define least-privilege access, role-based approvals, authentication standards, privileged access controls, and acceptable use expectations.",
        "Formal access policy with enough detail to ground the demo.",
        ["IT_Support_and_Access.md", "IT_Support.txt", "IT_Onboarding_Guide.md"],
        [
            theme("Access Principles and Governance Model", "least privilege and system ownership", "treat access as something granted for a task, not a convenience to collect over time", "hybrid work and the handling of sensitive employee, customer, and finance data make visibility into permissions especially important", "use system owners and documented role profiles rather than ad hoc approvals", "a finance employee receives the standard reporting profile while broader export rights require additional review"),
            theme("Authentication, MFA, and Identity Assurance", "managed identity as the basis of secure access", "keep recovery paths current and report suspicious authentication behavior immediately", "remote access patterns make MFA reliability and disciplined recovery especially important", "never send employees toward shared or unofficial workarounds to unblock urgent work", "an employee receives repeated unrecognized MFA prompts and reports them before retrying sign-in"),
            theme("Role-Based Access, Elevated Permissions, and Temporary Rights", "using standard profiles and review dates", "request elevated access only when the task really requires it and include an expected duration", "temporary elevated rights become risky when they quietly persist across remote and hybrid work", "prefer narrow, time-bounded access over broad permanent admin rights", "a data analyst receives temporary export rights for an audit instead of full administrator access"),
            theme("Shared Accounts, Service Accounts, and Prohibited Practices", "non-personal accounts under stronger controls", "avoid shared credentials, copied passwords, and personal storage of secrets", "where sensitive California employee or customer data is involved, shared-access patterns need especially strong controls", "require documented ownership, custody, rotation, and review when any non-personal account exists", "a legacy vendor portal requires a controlled shared account with explicit ownership and rotation rules"),
            theme("Access Reviews, Offboarding, and Change Management", "permissions across the full employment lifecycle", "trigger mover and leaver workflows promptly and challenge stale access during reviews", "distributed teams make stale access harder to notice if no one reviews it deliberately", "coordinate manager, People Operations, IT, and system-owner actions so lifecycle timing stays accurate", "an employee transfers from support to finance and the old support admin rights are reviewed instead of left active"),
            theme("Acceptable Use, Data Handling, and Reporting Responsibilities", "the obligations that come with access", "use approved tools and report accidental exposure or unusual data movement as soon as it is noticed", "employee, customer, and finance data handled in California operations can create legal and trust consequences when moved through unapproved channels", "model secure behavior and escalate questionable shortcuts instead of silently tolerating them", "an employee exports a report for home analysis and uses approved storage rather than a personal account"),
        ],
        [
            faq("What does least privilege mean in practice?", "It means employees receive only the access needed to do current work, not every access that might be convenient later."),
            faq("Can a manager approve any access they want?", "No. Managers sponsor business need, but system owners and security controls still determine what can actually be granted."),
            faq("Are shared credentials allowed?", "Only in rare approved scenarios with documented controls. Ordinary employee accounts must remain personal and attributable."),
            faq("When should temporary access expire?", "As soon as the task ends or at the review date established during approval."),
            faq("Why are access reviews necessary?", "Because roles change, projects end, and stale permissions are hard to notice without an intentional cycle."),
            faq("What should I do if I think I have too much access?", "Report it. Over-provisioned access is a control issue even if you did nothing wrong to receive it."),
            faq("Does acceptable use cover home-office behavior?", "Yes. Company data handling and approved-tool rules apply regardless of location."),
            faq("Who handles accidental data exposure?", "Report it immediately to the approved IT or Security path so containment and review can begin."),
        ],
        ["Security Operations", "Identity governance owner", "IT Service Desk for routing and intake"],
    ),
])


def main() -> None:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    for spec in DOCUMENT_SPECS:
        text = render_text(spec)
        count = word_count(text)
        if count < MIN_WORDS:
            raise RuntimeError(f"{spec['filename']} is too short: {count} words")
        path = DOCUMENTS_DIR / spec["filename"]
        if path.suffix.lower() in {".md", ".txt"}:
            path.write_text(text, encoding="utf-8")
        else:
            write_docx(path, text)
        print(f"{spec['filename']}: {count} words")


if __name__ == "__main__":
    main()
